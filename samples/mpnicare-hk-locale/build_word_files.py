#!/usr/bin/env python3
"""Build Traditional Chinese Word files documenting the MPNicare Hong Kong overlay."""
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path("/workspace/samples/mpnicare-hk-locale/word")
ART = Path("/cursor/stores/self/artifacts/mpnicare-hk-review")
OUT.mkdir(parents=True, exist_ok=True)
ART.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(0x12, 0x35, 0x3A)
TEAL = RGBColor(0x0B, 0x6E, 0x6E)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None, east="Microsoft JhengHei"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color:
        run.font.color.rgb = color
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(r, qn("w:rFonts"))
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), east)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=NAVY, east="Microsoft JhengHei")
    return p


def para(doc, text, *, size=11, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def shade_header(cell, hex_color="12353A"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    from docx.oxml import parse_xml
    tcPr.append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>'
    ))


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(255, 255, 255))
        shade_header(hdr[i])
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    p = doc.add_paragraph()
    r = p.add_run("MPNicare ｜ 藥華醫藥香港")
    set_run_font(r, size=10, color=TEAL, bold=True)
    h = doc.add_paragraph()
    r = h.add_run(title)
    set_run_font(r, size=20, bold=True, color=NAVY)
    para(doc, subtitle, size=11)
    para(doc, "日期：2026年8月21日  ·  敏感度：低（公開病人教育網站）  ·  狀態：實作規格／法律審閱草稿", size=10)
    return doc


def build_change_log():
    doc = setup_doc(
        "香港中文用語覆蓋層——改動清單",
        "版面、頁面結構及欄目位置維持不變。只改香港介面用語（選單、標題、頁尾、條款）及地區內容標示。三個地區欄目的文章正文不改寫。",
    )

    heading(doc, "1. 本檔用途", 1)
    para(doc, "本檔為 mpnicare.org「中文（香港）」語系的工作改動清單，供網頁（Wix 字串）、醫學事務（只審免責字句）及法務（使用條款頁）一併使用。內容反映 2026年8月21日審閱後的實作：地區提示不得出現在「病人教育專區」，只加在三個原文欄目。")

    heading(doc, "2. 標示規則（不得改寫文章）", 1)
    para(doc, "只適用於以下三欄（選單名稱會改；正文維持原文）：血腫專欄 → 血液腫瘤專欄；健康照護 → 健康資訊；品味生活（標題由「生活品味」統一）。")
    add_table(
        doc,
        ["位置", "加入以下字句", "不要加在"],
        [
            ["欄目頁——現有標題下、文章列表上", "其他地區內容提示\n本欄文章由香港以外地區的團隊或醫護專業人員提供，內容未必適用於香港，僅供一般參考。", "病人教育專區、關於MPN、首頁整頁、友善連結、使用條款"],
            ["文章卡——現有分類名稱旁", "其他地區內容", "病人教育專區的分類卡"],
            ["文章頁——標題或作者欄下一行", "其他地區內容｜原文保留，內容未必適用於香港。", "香港介面頁面"],
        ],
        [5.2, 7.5, 4.5],
    )
    para(doc, "首頁：不放整頁警告框。三張最新文章卡只在分類名稱旁加小標籤。")

    heading(doc, "3. 已實作的用語改動", 1)
    add_table(
        doc,
        ["頁面", "原文", "出現位置／原句", "香港用語", "備註"],
        [
            ["網站選單", "血腫專欄", "選單項目", "血液腫瘤專欄", "「血腫」≠血液腫瘤科。只改選單。"],
            ["網站選單", "健康照護", "選單項目", "健康資訊", "只改選單；文章正文不變。"],
            ["選單＋頁面標題", "品味生活 / 生活品味", "選單與標題用字次序不一", "品味生活（兩者統一）", "統一即可；不要改寫文章。"],
            ["頁尾（全站）", "信息", "非本網站所有之信息", "資料", "香港共用頁尾。"],
            ["頁尾（全站）", "帳號", "加入MPN官方line帳號", "賬號", "香港用「賬」。全站頁尾，故須改。"],
            ["頁尾（全站）", "智慧財產權", "為了尊重授權使用之智慧財產權", "知識產權", "香港全站頁尾。"],
            ["首頁", "並發症", "易引發嚴重並發症甚至危及生命", "併發症", "醫學用字應為「併」。現有 /zh-hk 誤用「並」。"],
            ["關於MPN", "紅細胞", "真性紅細胞增多症 / 主要是紅細胞增生", "紅血球", "避免機器翻譯回退。"],
            ["關於MPN", "白細胞", "血小板與白細胞增多", "白血球", "與「紅血球」對齊。"],
            ["關於MPN", "血細胞", "若同頁出現", "血球", "血球不要寫成「⋯細胞」。"],
            ["關於MPN", "按我了解更多", "按我了解更多MPN", "按此了解更多", "香港網頁少用「按我」。"],
            ["教育專區頁", "衛教專區", "頁面／分類名稱", "病人教育專區", "只改名稱。此頁不加地區橫幅。"],
            ["教育專區頁", "紅細胞增多症", "認識紅細胞增多症", "認識紅血球增多症", "標題對齊。"],
            ["教育專區頁", "有愛相髓病友故事集", "分類標題", "病友故事集", "只改名稱；故事正文不變。"],
            ["使用條款", "無限製", "您無限製或無條件地接受這些條款", "無限制", "簡轉繁錯誤。"],
            ["使用條款", "個資", "刪除您的個資等權利", "個人資料", "香港法律文本不用「個資」。"],
            ["使用條款", "個人資料保護法 第3、10、11條", "台灣個資法條號", "按《個人資料（私隱）條例》（第486章）重寫", "不要保留第3／10／11條。見檔案二。"],
            ["使用條款", "互聯網事業", "如同非互聯網事業的使用方式", "網上業務", "香港商業用語。"],
            ["使用條款", "藥華藥 / 藥華醫藥", "兩個簡稱混用", "藥華醫藥（首次：PharmaEssentia 藥華醫藥）", "統一。"],
            ["使用條款", "百份百 / 資訊", "無法保證資訊百份百正確無誤", "百分百 / 資料", "改正用字，並與頁尾「資料」對齊。"],
        ],
        [2.8, 3.2, 4.2, 4.2, 3.8],
    )

    heading(doc, "4. 不得改動的項目", 1)
    para(doc, "• 血腫專欄／健康照護／品味生活的全部文章正文（即使選單已改名）。")
    para(doc, "• 該等文章內的治療方法、藥名、研究、劑量及百分比。")
    para(doc, "• 該等文章內的本地用語、醫院名稱、產品名稱及系統名稱。")
    para(doc, "• 現有版面、分頁及欄目結構。")
    para(doc, "• 友善連結頁：無須改字。現有非香港機構連結保留；日後再以額外列加入香港機構。")

    heading(doc, "5. 不要原樣沿用現有 Wix /zh-hk 語系包", 1)
    para(doc, "現時公開頁 https://www.mpnicare.org/zh-hk 會清空三個文章欄、引入「紅細胞／白細胞／並發症／無限製」，並仍引用台灣《個人資料保護法》第3、10、11條。應改用本覆蓋層：保留原文＋加上標示＋改香港介面用語。")

    heading(doc, "6. 審閱負責人", 1)
    add_table(
        doc,
        ["項目", "負責人", "待決事項"],
        [
            ["Wix 字串表及三處標示位置", "產品經理＋網頁", "實作覆蓋層，不要重建版面範本"],
            ["三句免責字句（只審這三句）", "醫學事務", "確認用字；不要改寫文章"],
            ["第486章條款（檔案二）", "法務／合規", "上線前簽署確認"],
            ["頁面用語及標示的《不良廣告（醫藥）條例》核對", "產品經理＋法務", "待用戶提供第231章檔案後再核"],
        ],
        [5.5, 4.0, 7.5],
    )
    para(doc, "本站為病人教育網站，並非 BESREMi 推廣物料。不加新產品聲稱。原文文章中的 ET／MF 內容，一律標示為非香港內容。")

    path = OUT / "MPNicare-HK-wording-change-log.docx"
    doc.save(path)
    return path


def build_terms():
    doc = setup_doc(
        "使用條款與私隱政策——香港重寫稿（法律審閱草稿）",
        "不是把台灣《個人資料保護法》第3、10、11條逐字替換。須按香港法例第486章重寫。法務簽署前不得上線。",
    )

    heading(doc, "1. 現有頁面對照", 1)
    add_table(
        doc,
        ["現有用字（不要保留）", "香港用字"],
        [
            ["無限製", "無限制"],
            ["個資", "個人資料"],
            ["個人資料保護法第三條／第10條／第11條", "《個人資料（私隱）條例》（第486章）——查閱資料第18條；改正第22條；拒絕第20／24條；費用第28條"],
            ["互聯網事業", "網上業務"],
            ["「藥華藥」與「藥華醫藥」混用", "首次寫 PharmaEssentia 藥華醫藥；其後一律「藥華醫藥」"],
            ["資訊百份百", "資料百分百"],
            ["使用條款與隱私權政策", "使用條款與私隱政策"],
        ],
        [8.0, 9.0],
    )

    heading(doc, "2. 香港頁面草稿全文", 1)
    para(doc, "使用條款與私隱政策", size=14, bold=True)

    blocks = [
        "本網站由 PharmaEssentia 藥華醫藥股份有限公司（以下簡稱「藥華醫藥」）所建立經營。我們重視每一位使用者所享有的服務，特此說明本網站的使用政策，以保障您的權益，請您細讀本使用條款的內容。",
        "一般性資料",
        "對於本站所載資料的取覽或使用，您受下列條款和條件以及所有適用法律所規範。經由取覽或瀏覽本網站，您無限制或無條件地接受這些條款和條件，並確認其將取代您和藥華醫藥之間其他任何的協議。",
        "醫療資料",
        "本網站之醫療相關資料僅供科學資料或病人教育目的使用。由於醫療科技的發展日新月異，本公司無法保證資料百分百正確無誤。任何疾病治療、醫學問題及專業知識，應諮詢您的專業醫生及護士，本站不提供任何診斷、用藥或治療建議。",
        "私隱政策（香港）",
        "藥華醫藥尊重訪客私隱，並按香港法例第486章《個人資料（私隱）條例》（下稱「《條例》」）的保障資料原則處理個人資料。",
        "您可向本公司提出：（一）查閱資料要求（《條例》第18條）；以及（二）改正資料要求（《條例》第22條）。",
        "本公司會在《條例》規定的時限內回覆。在第20條所列明的情況下，本公司可拒絕依從查閱資料要求；在第24條所列明的情況下，可拒絕依從改正資料要求。本公司亦可按第28條收取不超逾直接成本的合理費用。",
        "《條例》並無賦予與台灣《個人資料保護法》第3條相同的概括「刪除／停止處理」法定權利。本公司仍會按保障資料第2原則（保留期間）及第3原則（使用限制），在不再需要時刪除或停止使用個人資料。您亦可要求退出通訊服務。是否接納個別申請，須視乎執行業務所必須、法定保存期間及《條例》准許的豁免而定。",
        "個人資料的使用",
        "個人資料包括您的姓名、地址、電話號碼、電郵地址，或其他可合理用於識別您的資料。藥華醫藥只在您自願提供時收集該等資料。",
        "當藥華醫藥收到個人資料時，我們可能將資料用於合理商業用途，如同非網上業務的使用方式。例如，我們可能透過通訊軟件或定期訊息聯絡您。為提供準確服務，收集的資料亦可能用於編製統計及分析，但不會公開可識別個別人士的資料。",
        "您的資料可能傳送至位於其他國家的藥華醫藥附屬機構處理。藥華醫藥不會出售或出租個人資料，除非事先通知並取得您的明確同意。",
        "請聯絡我們：sales@museshc.com",
    ]
    for b in blocks:
        is_h = b in {"一般性資料", "醫療資料", "私隱政策（香港）", "個人資料的使用"}
        para(doc, b, size=12 if is_h else 11, bold=is_h)

    heading(doc, "3. 法務審閱要點", 1)
    para(doc, "• 核對查閱／改正／拒絕／費用條文，是否符合現行第486章及個人資料私隱專員公署指引。")
    para(doc, "• 確認自願刪除／退出通訊的字句，應否保留為公司政策（而非台灣個資法第3條的法定權利）。")
    para(doc, "• 確認藥華醫藥附屬機構跨境傳送個人資料的用字。")
    para(doc, "• 《不良廣告（醫藥）條例》（第231章）不在本檔範圍；待用戶上傳條例檔案後，另核介面用語及標示。")
    para(doc, "• 法務簽署前，不得將本草稿作為正式私隱聲明上線。")

    path = OUT / "MPNicare-HK-terms-privacy-rewrite.docx"
    doc.save(path)
    return path


def main():
    a = build_change_log()
    b = build_terms()
    for p in (a, b):
        dest = ART / p.name
        dest.write_bytes(p.read_bytes())
        print("wrote", p, p.stat().st_size)
        print("artifact", dest)


if __name__ == "__main__":
    main()
