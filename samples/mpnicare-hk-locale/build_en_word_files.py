#!/usr/bin/env python3
"""Build Traditional Chinese + English Word files for the MPNicare English overlay."""
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path("/workspace/samples/mpnicare-hk-locale/word")
ART = Path("/cursor/stores/self/artifacts/mpnicare-hk-review")
OUT.mkdir(parents=True, exist_ok=True)
ART.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(0x12, 0x35, 0x3A)
TEAL = RGBColor(0x0B, 0x6E, 0x6E)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None, east="Microsoft JhengHei"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color:
        run.font.color.rgb = color
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(r, qn("w:rFonts"))
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), east)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=NAVY)
    return p


def para(doc, text, *, size=11, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def shade_header(cell, hex_color="12353A"):
    tcPr = cell._tc.get_or_add_tcPr()
    from docx.oxml import parse_xml
    tcPr.append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{hex_color}"/>'
    ))


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = ""
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(255, 255, 255))
        shade_header(table.rows[0].cells[i])
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def setup_doc(title, subtitle):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(1.8)
    s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.0)
    s.right_margin = Cm(2.0)
    p = doc.add_paragraph()
    r = p.add_run("MPNicare ｜ 藥華醫藥香港")
    set_run_font(r, size=10, color=TEAL, bold=True)
    h = doc.add_paragraph()
    r = h.add_run(title)
    set_run_font(r, size=20, bold=True, color=NAVY)
    para(doc, subtitle, size=11)
    para(doc, "日期：2026年9月2日  ·  敏感度：低（公開病人教育網站）  ·  狀態：建議用語／法律審閱草稿", size=10)
    return doc


def build_change_log():
    doc = setup_doc(
        "英文版建議改動清單（取代 Google 翻譯）",
        "現有 https://www.mpnicare.org/en 為機器翻譯。建議與香港中文版同一做法：只改介面用語及條款，三欄文章正文維持原文並加上地區標示。不要沿用空白欄目。",
    )

    heading(doc, "1. 建議做法", 1)
    para(doc, "不要把現有 /en Google 翻譯原樣上線。機器翻譯已出現臨床錯誤（把「原發性血小板過多症」譯成 thrombocytopenia「血小板過低」；把「病友故事」譯成 myelopathy「脊髓病」），亦把「血腫專欄」譯成 Hematoma（血腫／瘀傷）。")
    para(doc, "建議：用人工修訂的英文介面用語覆蓋現有字串；文章欄顯示中文原文＋「Content from another region」標示。日後如需英文全文，只挑重點文章交專業醫學翻譯，不要整站機翻。")

    heading(doc, "2. 選單必須與中文版同一結構", 1)
    para(doc, "中文版頂層只有 7 項：首頁、關於MPN▾、衛教專區▾、血腫專欄、健康照護、品味生活、友善連結。英文版多了「新網頁」（報名系統）、兩個「關於」、搜尋結果，標籤又較長，Wix 才出現 More，More 裡再堆兩個中文「關於」。")
    add_table(
        doc,
        ["中文版（準則）", "現有 /en（不要保留）", "建議英文"],
        [
            ["首頁", "front page", "Home"],
            ["（無此項）", "新網頁 → /en/報名系統", "刪除，不要出現在選單"],
            ["關於MPN▾（5 個中文下拉）", "關於MPN（下拉仍全中文）", "About MPN▾ + 5 個英文下拉"],
            ["衛教專區▾（4 個中文下拉）", "關於 + 副本 副本 衛教手冊…", "Patient education▾ + 4 個英文下拉"],
            ["血腫專欄", "Hematoma column", "Hematology"],
            ["健康照護", "health care", "Health information"],
            ["品味生活", "Lifestyle", "Lifestyle"],
            ["友善連結", "關於（連到 about-4）", "Useful links"],
            ["（條款只在頁尾）", "關於（連到 about-5 條款）", "刪除選單項；條款留在頁尾"],
            ["（無搜尋項）", "搜尋結果 + More", "刪除"],
        ],
        [5.5, 6.5, 5.0],
    )
    para(doc, "Wix 操作：英文語系隱藏未對應中文選單的頁；關閉自動加入新頁；選單列不要用 More 溢出。頁尾 LINE 與私隱／Cookies 分兩欄，不要用會重疊的絕對定位。")

    heading(doc, "3. 必須改掉的錯誤（臨床／品牌）", 1)
    add_table(
        doc,
        ["頁面", "現有 /en（Google 翻譯）", "建議英文", "為何要改"],
        [
            ["選單", "Hematoma column", "Hematology", "hematoma＝血腫／瘀傷，不是血液腫瘤。"],
            ["選單", "front page", "Home", "網站選單應用 Home。"],
            ["選單", "health care", "Health information", "與香港中文「健康資訊」對齊；較「health care」準確。"],
            ["選單", "Lifestyle（可保留）", "Lifestyle", "與「品味生活」對齊即可。"],
            ["選單", "友善連結仍多為中文／Friendly Links", "Useful links", "Friendly links 是中式英語。"],
            ["教育專區", "Myeloproliferative tumor", "Myeloproliferative neoplasms", "WHO／ASH 用 neoplasm，不用 tumor。"],
            ["教育專區", "Understanding Polycythemia", "Understanding polycythemia vera", "病名要寫完整 PV。"],
            ["教育專區", "Understanding thrombocytopenia", "Understanding essential thrombocythemia", "thrombocytopenia＝血小板過低，與 ET 相反，屬臨床錯誤。"],
            ["教育專區", "A Collection of Stories from Patients with Myelopathy", "Patient stories", "myelopathy＝脊髓病，與骨髓無關。"],
            ["教育專區", "衛教專區／副本 副本 衛教手冊", "Patient education", "清掉 Wix「副本」頁名；改病人教育。"],
            ["關於MPN", "myeloproliferative tumor", "myeloproliferative neoplasm", "同上。"],
            ["關於MPN", "Erythrocytosis vera (PV)", "Polycythemia vera (PV)", "國際通用病名是 polycythemia vera。"],
            ["關於MPN", "red blood cell hyperplasia / platelet and leukocytosis", "overproduction of red blood cells, often with elevated platelets and white blood cells", "現句文法及醫學用字都不對。"],
            ["關於MPN", "Essential thrombocytosis", "Essential thrombocythemia (ET)", "ASH／病友組織多用 thrombocythemia。"],
            ["關於MPN", "early / obvious stage of myelofibrosis", "prefibrotic PMF (pre-PMF) and overt PMF", "對齊國際分期用字。"],
            ["關於MPN", "Click me to learn more about MPN", "Learn more about MPN", "不要直譯「按我」。"],
            ["條款", "Medical sexual information", "Medical information", "把「醫療性資訊」的「性」誤譯成 sexual。"],
            ["條款", "PharmaEssence / PharmaEssential / PharmaEs / PharmaWorld / Yaohua Pharmaceutical", "PharmaEssentia（首次可加 藥華醫藥）", "公司名被機翻拆壞。"],
            ["條款", "Article 3 / 10 / 11 of the Personal Data Protection Act", "按香港《個人資料（私隱）條例》（第486章）重寫", "與香港中文版同一法律問題。"],
            ["三欄文章", "No posts published in this language yet", "顯示原文＋地區標示", "不要清空欄目。"],
            ["頂層選單", "新網頁（連到報名系統）", "刪除。中文版沒有此項", "Wix 多語系把未上架頁加進英文選單。"],
            ["頂層選單", "兩個額外「關於」（友善連結／條款）", "刪除。條款只放頁尾；友善連結用 Useful links", "More 裡出現兩個關於，即因此。"],
            ["頂層選單", "搜尋結果 + More", "刪除。中文版沒有搜尋項；More 是溢出選單", "英文標籤較長，Wix 把多出的項推進 More。"],
            ["關於MPN 下拉", "血液數據怎麼看／什麼是MPN？／真性紅血球增生症…（全中文）", "How to read blood results / What is MPN? / Polycythemia vera (PV) / Essential thrombocythemia (ET) / Primary myelofibrosis (PMF)", "下拉必須與中文版同一 5 項，並譯成英文。"],
            ["衛教下拉（現標成「關於」）", "副本 副本 衛教手冊／副本 衛教手冊／副本 認識紅血球增多症", "Patient education▾：Myeloproliferative neoplasms / Understanding polycythemia vera / Understanding essential thrombocythemia / Patient stories", "清掉 Wix「副本」頁名；不要把衛教專區譯成 About。"],
            ["頁尾排版", "LINE 歡迎句與 cookies／私隱政策重疊", "LINE 改短句並獨立一欄：Official MPN LINE / Latest updates", "英文比中文長，Wix 絕對定位會撞字。"],
        ],
        [2.6, 4.6, 4.4, 5.4],
    )

    heading(doc, "4. 首頁建議英文（可直接貼上 Wix）", 1)
    para(doc, "This is a patient education website on myeloproliferative neoplasms (MPN).", bold=True)
    para(doc, "MPN is a bone marrow disease caused by genetic mutations in hematopoietic stem cells, leading to abnormal overproduction of blood cells. Symptoms are often nonspecific, but over time the disease can cause serious — even life-threatening — complications. For this reason it is sometimes called a silent killer.")
    para(doc, "Here you can find information about MPN, treatment updates, and care guidance to help you understand the condition more clearly.")
    para(doc, "現有英文單張申請句已通順，建議保留：An English version of the patient education leaflet is available. Please complete the form below, and we will send you a digital copy via email.")

    heading(doc, "5. 關於MPN 建議英文", 1)
    para(doc, "What is MPN?")
    para(doc, "MPN stands for myeloproliferative neoplasm. It is a blood cancer in which abnormal hematopoietic stem cells in the bone marrow cause abnormal blood counts. Most patients have a known driver mutation, such as JAK2, CALR, or MPL. The three most common MPNs are:")
    para(doc, "1. Polycythemia vera (PV): mainly overproduction of red blood cells, often with elevated platelets and white blood cells.")
    para(doc, "2. Essential thrombocythemia (ET): mainly overproduction of platelets.")
    para(doc, "3. Primary myelofibrosis (PMF): overgrowth of fibroblasts in the bone marrow, which prevents normal blood-cell production. It is divided into prefibrotic PMF (pre-PMF) and overt PMF.")
    para(doc, "Learn more about MPN")

    heading(doc, "6. 頁尾／選單建議", 1)
    add_table(
        doc,
        ["位置", "現有 /en", "建議"],
        [
            ["最新文章", "Latest articles", "Latest articles（可保留）"],
            ["LINE", "Join the official MPN Line account / We will provide you with the latest information.（與私隱政策重疊）", "Official MPN LINE / Latest updates（獨立右欄，不要蓋住條款）"],
            ["Cookies", "…cookie policy terms set out on our legal and privacy pages.", "This website uses cookies to provide a better experience. By continuing, you agree to our cookie policy."],
            ["友善連結機構名", "Taiwan Myeloproliferative Tumor Care Association", "Taiwan Myeloproliferative Neoplasm Care Association（或保留該會官方英文名）"],
        ],
        [4.0, 7.0, 6.0],
    )

    heading(doc, "7. 不要改的範圍", 1)
    para(doc, "• 三欄文章正文、藥名、劑量、研究數據、醫院及產品名稱：不要用 Google 翻譯改寫。")
    para(doc, "• 版面、分欄、顏色、LINE 頁尾位置。")
    para(doc, "• 病人教育專區：只改分類名稱，不加地區橫幅（與香港中文版相同）。")
    para(doc, "• 不加 BESREMi 或其他新產品聲稱。")

    heading(doc, "8. 審閱負責人", 1)
    add_table(
        doc,
        ["項目", "負責人", "待決"],
        [
            ["英文介面字串表", "產品經理＋網頁", "覆蓋 /en，不要重建版面"],
            ["病名（PV／ET／PMF／MPN）", "醫學事務", "確認用 WHO／ASH 英文"],
            ["第486章英文條款", "法務／合規", "上線前簽署；勿保留台灣個資法條號"],
            ["是否另做重點文章專業翻譯", "產品經理＋醫學事務", "非必須；現階段用原文＋標示即可"],
        ],
        [5.5, 4.5, 7.0],
    )

    path = OUT / "MPNicare-EN-wording-change-log.docx"
    doc.save(path)
    return path


def build_terms():
    doc = setup_doc(
        "Terms of use and privacy policy — English rewrite (legal review draft)",
        "Replace the live /en terms page. Do not keep Taiwan PDPA Arts. 3, 10 and 11, “Medical sexual information”, or broken company names. Rewrite under Hong Kong Cap. 486. Do not publish until Legal signs off.",
    )
    heading(doc, "1. 現有 /en 必須刪除", 1)
    add_table(
        doc,
        ["現有 Google 翻譯", "建議"],
        [
            ["Medical sexual information", "Medical information"],
            ["PharmaEssence / PharmaEssential / PharmaEs / PharmaWorld / Yaohua Pharmaceutical / PharmaEssentials", "PharmaEssentia"],
            ["Articles 3, 10 and 11 of the Personal Data Protection Act", "Personal Data (Privacy) Ordinance (Cap. 486) — access s.18; correction s.22; refusal s.20 / s.24; fee s.28"],
            ["Collective/cluster data", "可簡寫為 aggregated / cookie data，或併入 cookies 段"],
        ],
        [8.5, 8.5],
    )
    heading(doc, "2. Draft English page", 1)
    blocks = [
        "This website is operated by PharmaEssentia Corporation (referred to below as “PharmaEssentia”). Please read these terms carefully. They explain how you may use the site and how we handle personal data.",
        "General information",
        "Your access to or use of this website is governed by these terms and all applicable laws. By browsing the site, you accept these terms without limitation. They replace any earlier agreement between you and PharmaEssentia about this website.",
        "Medical information",
        "Medical content on this website is for scientific information and patient education only. Medical science changes quickly, so we cannot guarantee that every item is complete or current. For diagnosis, medication, or treatment, please consult your doctor or nurse. This website does not give medical advice.",
        "Privacy policy (Hong Kong)",
        "PharmaEssentia respects your privacy and handles personal data in line with the data protection principles of the Personal Data (Privacy) Ordinance (Cap. 486) (the “Ordinance”).",
        "You may make a data access request (section 18) and a data correction request (section 22). We will reply within the time allowed by the Ordinance. We may refuse a request in the cases listed in sections 20 and 24, and may charge a fee that does not exceed direct cost (section 28).",
        "The Ordinance does not give the same general statutory right to deletion as Article 3 of Taiwan’s Personal Data Protection Act. We still delete or stop using personal data when we no longer need it. You may also opt out of communications.",
        "How we use personal data",
        "We collect your name, address, telephone number, email address, or other identifying information only when you choose to give it. We may use it for reasonable business purposes, including contact by messaging apps or updates, and for statistics that do not identify you. Data may be sent to PharmaEssentia affiliates in other countries. We will not sell or rent personal data unless we tell you first and you give clear consent.",
        "Please contact us: sales@museshc.com",
    ]
    heads = {"General information", "Medical information", "Privacy policy (Hong Kong)", "How we use personal data"}
    for b in blocks:
        para(doc, b, size=12 if b in heads else 11, bold=b in heads)
    heading(doc, "3. 法務要點", 1)
    para(doc, "• 核對第486章查閱／改正／拒絕／費用條文及私隱專員公署指引。")
    para(doc, "• 公司英文法定名稱以公司註冊為準（PharmaEssentia Corporation / PharmaEssentia 藥華醫藥）。")
    para(doc, "• 法務簽署前不得上線。")
    path = OUT / "MPNicare-EN-terms-privacy-rewrite.docx"
    doc.save(path)
    return path


def main():
    for p in (build_change_log(), build_terms()):
        dest = ART / p.name
        dest.write_bytes(p.read_bytes())
        print("wrote", p, p.stat().st_size)
        print("artifact", dest)


if __name__ == "__main__":
    main()
