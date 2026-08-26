#!/usr/bin/env python3
"""Convert a subset of Markdown (headings, tables, lists, quotes, bold) to .docx."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


BANNER = (
    "INTERNAL REHEARSAL ONLY — not an approved leave-behind, detail aid, "
    "or patient material. Any HCP-facing use requires Medical + Compliance "
    "+ global lexicon clearance."
)


def set_run_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_formatted_text(paragraph, text: str, size=11, default_bold=False, default_italic=False):
    """Split **bold** and *italic* / `code` into runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, size=size, bold=default_bold, italic=default_italic)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, italic=default_italic)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=size - 1, italic=default_italic)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, bold=default_bold, italic=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=default_bold, italic=default_italic)


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text: str, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    add_formatted_text(p, text.strip(), size=size, default_bold=bold)


def is_table_separator(line: str) -> bool:
    s = line.strip()
    return bool(re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", s))


def parse_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            val = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            set_cell_text(cell, val, bold=(i == 0), size=9 if i else 10)
            if i == 0:
                shade_cell(cell, "1F4E79")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True
            elif i % 2 == 0:
                shade_cell(cell, "F2F2F2")
    doc.add_paragraph()


def convert(md_path: Path, docx_path: Path, title: str | None = None):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "BESREMi HK PM — internal training  |  Not for HCP leave-behind"
    set_run_font(header.runs[0], size=8, italic=True, color=RGBColor(127, 127, 127))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    set_run_font(run, size=8, color=RGBColor(127, 127, 127))
    # simple page field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    r = footer.add_run()._element
    r.append(fldChar1)
    r.append(instr)
    r.append(fldChar2)

    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = banner.add_run(BANNER)
    set_run_font(run, size=10, bold=True, color=RGBColor(192, 0, 0))
    banner.paragraph_format.space_after = Pt(12)

    if title:
        p = doc.add_heading(title, level=0)
        for run in p.runs:
            run.font.color.rgb = RGBColor(31, 78, 121)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and is_table_separator(lines[i + 1]):
            rows = [parse_table_row(stripped)]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                if not is_table_separator(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 3)
            heading = doc.add_heading("", level=level)
            add_formatted_text(heading, m.group(2).strip(), size=16 - level * 2, default_bold=True)
            i += 1
            continue

        # blockquote (possibly multi-line)
        if stripped.startswith(">"):
            chunks = []
            while i < n and lines[i].strip().startswith(">"):
                q = re.sub(r"^>\s?", "", lines[i].strip())
                chunks.append(q)
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(10)
            add_formatted_text(p, " ".join(chunks), size=11, default_italic=True)
            for run in p.runs:
                if run.font.color.rgb is None:
                    run.font.color.rgb = RGBColor(89, 89, 89)
            continue

        # unordered list
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # numbered list
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main():
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "training-pack" / "word"
    jobs = [
        (
            root / "playbooks" / "physician-product-training.md",
            out_dir / "01-BESREMi-HK-PM-physician-product-training.docx",
            "BESREMi HK PM — Physician product training",
        ),
        (
            root / "industry" / "pm-reference" / "pegasys-claim-check.md",
            out_dir / "02-BESREMi-HK-PM-Pegasys-claim-check.docx",
            "BESREMi HK PM — Pegasys claim check",
        ),
        (
            root / "industry" / "pm-reference" / "field-visit-patterns.md",
            out_dir / "03-BESREMi-HK-PM-field-visit-patterns-deidentified.docx",
            "BESREMi HK PM — Field visit patterns (de-identified)",
        ),
    ]
    for src, dest, title in jobs:
        convert(src, dest, title=title)


if __name__ == "__main__":
    sys.exit(main() or 0)
